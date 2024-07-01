import configparser
import pandas as pd
from datetime import datetime, timedelta

from PyQt5.QtCore import QDateTime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from sqlalchemy import create_engine
import matplotlib.pyplot as plt

from validator.user_manager import userManager

def load_config():
    config = configparser.ConfigParser()
    config.read('config.ini')
    return config

def save_config(report_path):
    config = configparser.ConfigParser()
    config.read('config.ini')

    if report_path:
        if 'DEFAULT' not in config:
            config['DEFAULT'] = {}
        config['DEFAULT']['report_path'] = report_path

    with open('config.ini', 'w') as configfile:
        config.write(configfile)

# Replace 'username', 'password', 'localhost', 'dbname' with your actual MySQL credentials and database name
engine = create_engine('mysql+pymysql://root:root@localhost/poswithinventorysystem')

# Query to join order, leftover, and package tables
query = """
SELECT 
    o.Order_ID, 
    o.Date,
    o.Time,
    o.Total_Amount,
    o.Payment_Status,
    o.Package_ID,
    o.Leftover_ID,
    o.Customer_Name,
    o.Soup_Variation,
    o.Guest_Pax,
    o.Time_Status,
    o.Order_Type,
    o.Payment_Method,
    o.Cash_Amount,
    o.reference_id,
    o.Discount_Type,
    o.Priority_Order,
    o.Subtotal_Amount,
    o.VAT_Amount,
    o.Discount_Amount,
    o.Change_Amount,
    o.Package_Total_Amount,
    o.Add_Ons_Total_Amount,
    l.Grams as Leftover_Grams,
    l.Penalty_Fee,
    p.Package_Name,
    p.Package_Price
FROM 
    `order` o
LEFT JOIN 
    leftover l ON o.Leftover_ID = l.Leftover_ID
LEFT JOIN 
    package p ON o.Package_ID = p.Package_ID
"""

dataframe = pd.read_sql(query, engine)

# Combine Date and Time into a single DateTime column
dataframe['DateTime'] = dataframe['Date'].astype(str) + ' ' + dataframe['Time'].astype(str)
dataframe['DateTime'] = pd.to_datetime(dataframe['DateTime'])

# Function to generate a daily report
def generate_daily_report():
    data = dataframe.copy()
    today = datetime.now().date()
    daily_data = data[data['DateTime'].dt.date == today]
    return daily_data

# Function to generate a weekly report
def generate_weekly_report():
    data = dataframe.copy()
    today = datetime.now().date()
    last_week = today - timedelta(days=7)
    weekly_data = data[(data['DateTime'].dt.date >= last_week) & (data['DateTime'].dt.date <= today)]
    return weekly_data

# Function to generate a monthly report
def generate_monthly_report():
    data = dataframe.copy()
    today = datetime.now().date()
    last_month = today - timedelta(days=30)
    monthly_data = data[(data['DateTime'].dt.date >= last_month) & (data['DateTime'].dt.date <= today)]
    return monthly_data

# Function to generate plots
def plot_reports(report_data, frequency, file_path):
    # Convert DateTime to string format for plotting
    report_data['DateTime_str'] = report_data['DateTime'].dt.strftime('%B %d, %Y')
    report_data['Time'] = report_data['DateTime'].dt.strftime('%I:%M %p')

    today = datetime.today()
    date_today = today.strftime("%B %d, %Y")
    last_month = today - timedelta(days=30)
    last_month_range = f'{last_month.strftime("%B %d, %Y")} to {today.strftime("%B %d, %Y")}'
    last_week = today - timedelta(days=7)
    last_week_range = f'{last_week.strftime("%B %d, %Y")} to {today.strftime("%B %d, %Y")}'

    # Sort report data by DateTime for chronological order
    sorted_report_data = report_data.sort_values(by='Total_Amount', ascending=True)

    # Group 'x' data by weeks within the month
    weekly_dates = sorted_report_data.groupby(pd.Grouper(key='DateTime', freq='W'))['DateTime_str'].first()

    # Group 'y' data by weeks within the month
    sorted_report_data['Total_Amount'] = sorted_report_data['Total_Amount'].astype(float)
    weekly_total_amount = sorted_report_data.groupby(pd.Grouper(key='DateTime', freq='W'))['Total_Amount'].sum()

    # Total Sales
    plt.figure(figsize=(10, 6))
    plt.xlabel('Date')
    plt.ylabel('Total Amount')
    if frequency == 'Daily':
        plt.plot(report_data['Time'], sorted_report_data['Total_Amount'], color='blue', marker='o')
        plt.title(f'{frequency} Total Sales ({date_today})')
    elif frequency == 'Weekly':
        plt.plot(report_data['DateTime_str'], sorted_report_data['Total_Amount'], color='blue', marker='o')
        plt.title(f'{frequency} Total Sales ({last_week_range})')
    elif frequency == 'Monthly':
        plt.plot(weekly_dates, weekly_total_amount, color='blue', marker='o')
        plt.title(f'{frequency} Total Sales ({last_month_range})')

    plt.xticks(rotation=45)
    plt.grid(True)
    plt.tight_layout()
    plt.savefig(f'{file_path}/total_sales_{frequency.lower()}.png')
    plt.close()

    # Sales by Payment Method
    payment_counts = report_data['Payment_Method'].value_counts()
    plt.figure(figsize=(8, 6))
    plt.pie(payment_counts, labels=payment_counts.index,
            autopct=lambda pct: f'{pct:.1f}% ({int(pct / 100 * sum(payment_counts))})', startangle=140)
    if frequency == 'Daily':
        plt.title(f'{frequency} Sales by Payment Method ({date_today})')
    elif frequency == 'Weekly':
        plt.title(f'{frequency} Sales by Payment Method ({last_week_range})')
    elif frequency == 'Monthly':
        plt.title(f'{frequency} Sales by Payment Method ({last_month_range})')
    plt.legend(loc='best')
    plt.tight_layout()
    plt.savefig(f'{file_path}/sales_payment_method_{frequency.lower()}.png')
    plt.close()

    # Sales by Category
    category_counts = report_data['Order_Type'].value_counts()
    plt.figure(figsize=(8, 6))
    plt.pie(category_counts, labels=category_counts.index,
            autopct=lambda pct: f'{pct:.1f}% ({int(pct / 100 * sum(category_counts))})', startangle=140)
    if frequency == 'Daily':
        plt.title(f'{frequency} Sales by Category ({date_today})')
    elif frequency == 'Weekly':
        plt.title(f'{frequency} Sales by Category ({last_week_range})')
    elif frequency == 'Monthly':
        plt.title(f'{frequency} Sales by Category ({last_month_range})')
    plt.legend(loc='best')
    plt.tight_layout()
    plt.savefig(f'{file_path}/sales_category_{frequency.lower()}.png')
    plt.close()


def save_report_to_excel(report_data, report_type, file_path):
    filename = f'{file_path}/{report_type}_report.xlsx'
    report_data.to_excel(filename, index=False)
    print(f"{report_type.capitalize()} report has been generated and saved to '{filename}'")


# Function to save report to Word document
def save_report_to_word(report_data, report_type, file_path):
    document = Document()

    section = document.sections[0]
    header = section.header
    month = datetime.now().strftime("%B")
    content_header = [
        "Moon Hey Hotpot and Grill",
        "848A Banawe St, Quezon City, 1114 Metro Manila",
        "0917 624 9289",
        f"Sales {report_type} Report",
        f"({month})",
    ]

    for content_h in content_header:
        header_paragraph = header.add_paragraph(content_h)
        run = header_paragraph.runs[0]

        # Center align the paragraph
        header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Remove spacing before and after the paragraph
        header_paragraph.paragraph_format.space_before = Pt(0)
        header_paragraph.paragraph_format.space_after = Pt(0)

    footer = section.footer
    date = datetime.now().strftime("%B %d, %Y")
    time = datetime.now().strftime("%I:%M %p")
    user_manager = userManager._instance
    username = user_manager.get_current_username()
    footer_paragraph = footer.add_paragraph()
    footer_paragraph.text = f"{date} | {time}    Created by: {username}"
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Detailed sales analysis
    total_sales = report_data['Total_Amount'].astype(float).sum()
    total_transactions = report_data['Order_ID'].nunique()
    avg_transaction_value = total_sales / total_transactions if total_transactions != 0 else 0
    total_customers = report_data['Customer_Name'].nunique()
    payment_methods = report_data['Payment_Method'].value_counts(normalize=True) * 100
    sales_by_category = report_data['Order_Type'].value_counts()
    total_discounts = report_data['Discount_Amount'].astype(float).sum()
    total_add_ons = report_data['Add_Ons_Total_Amount'].astype(float).sum()
    payment_counts = report_data['Payment_Method'].value_counts()

    # Calculating profit
    report_data['Profit'] = report_data['Total_Amount'].astype(float) - report_data['Subtotal_Amount'].astype(float)
    total_profit = report_data['Profit'].sum()

    # Calculate sales from each package
    sales_by_package = report_data['Total_Amount'].astype(float).groupby(report_data['Package_Name']).sum()

    document.add_heading(f"Total Sales: ₱{total_sales:,.2f}", level=1)
    document.add_paragraph(f"Total Transactions: {total_transactions}")
    document.add_paragraph(f"Average Transaction Value: ₱{avg_transaction_value:,.2f}")
    document.add_paragraph(f"Total Customers Served: {total_customers}")
    document.add_paragraph(f"Total Discounts Given: ₱{total_discounts:,.2f}")
    document.add_paragraph(f"Total Profit: ₱{total_profit:,.2f}\n")

    document.add_heading("Sales by Payment Method", level=2)
    for method, percent in payment_methods.items():
        count = payment_counts[method]
        document.add_paragraph(f"{method}: {count} ({percent:.2f}%)")

    document.add_heading("Sales by Category", level=2)
    for category, count in sales_by_category.items():
        document.add_paragraph(f"{category}: {count}")

    document.add_heading("Sales by Add-Ons", level=2)
    document.add_paragraph(f"Total Add-Ons Sales: ₱{total_add_ons:,.2f}")

    document.add_heading("Sales by Package", level=2)
    for package, amount in sales_by_package.items():
        document.add_paragraph(f"{package}: ₱{float(amount):,.2f}")

    plot_paths = {
        'Total Sales': f'{file_path}/total_sales_{report_type.lower()}.png',
        'Sales by Payment Method': f'{file_path}/sales_payment_method_{report_type.lower()}.png',
        'Sales by Category': f'{file_path}/sales_category_{report_type.lower()}.png'
    }

    for title, path in plot_paths.items():
        document.add_heading(f"{title} ({report_type})", level=2)
        document.add_picture(path, width=Inches(6))

    filename = f'{file_path}/{report_type}_report.docx'
    document.save(filename)
    print(f"{report_type.capitalize()} report has been generated and saved to '{filename}'")

