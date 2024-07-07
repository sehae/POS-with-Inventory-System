import configparser
import json
import os

import pandas as pd
from datetime import datetime, timedelta
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx2pdf import convert
from sqlalchemy import create_engine
import matplotlib.pyplot as plt

from server.create_engine import get_db_engine
from validator.user_manager import userManager

def load_config():
    config = configparser.ConfigParser()
    config.read('config.ini')
    return config

def save_config(trend_path):
    config = configparser.ConfigParser()
    config.read('config.ini')

    if trend_path:
        if 'DEFAULT' not in config:
            config['DEFAULT'] = {}
        config['DEFAULT']['trend_path'] = trend_path

    with open('config.ini', 'w') as configfile:
        config.write(configfile)

engine = get_db_engine()

# Query to join order, leftover, package, and add_on tables
query = """
SELECT 
    o.Order_ID, 
    o.Date,
    o.Time,
    o.Total_Amount,
    o.Customer_Name,
    o.Soup_Variation,
    o.Guest_Pax,
    o.Order_Type,
    o.Payment_Method,
    p.Package_Name,
    a.Product_Details
FROM 
    `order` o
LEFT JOIN 
    leftover l ON o.Leftover_ID = l.Leftover_ID
LEFT JOIN 
    package p ON o.Package_ID = p.Package_ID
LEFT JOIN
    add_on a ON o.Order_ID = a.Order_ID
"""

dataframe = pd.read_sql(query, engine)

# Combine Date and Time into a single DateTime column
dataframe['DateTime'] = dataframe['Date'].astype(str) + ' ' + dataframe['Time'].astype(str)
dataframe['DateTime'] = pd.to_datetime(dataframe['DateTime'])

def load_product_data():
    # Query to get Product_Details and Order_ID from add_on table
    product_details_query = "SELECT Product_Details, Order_ID FROM add_on"
    product_details_data = pd.read_sql(product_details_query, engine)

    # Parse JSON and extract Product_ID and Quantity
    product_id_quantity_list = []
    for index, row in product_details_data.iterrows():
        product_details_json = json.loads(row['Product_Details'])
        for product in product_details_json:
            product_id_quantity_list.append((row['Order_ID'], product['product_id'], product['quantity']))

    # Convert list to DataFrame
    product_id_quantity_df = pd.DataFrame(product_id_quantity_list, columns=['Order_ID', 'Product_ID', 'Quantity'])

    # Join with product table to get Name
    product_query = "SELECT Product_ID, Name FROM product"
    product_data = pd.read_sql(product_query, engine)
    product_data = pd.merge(product_id_quantity_df, product_data, on='Product_ID')

    # Join with order table to get DateTime
    order_query = "SELECT Order_ID, Date, Time FROM `order`"
    order_data = pd.read_sql(order_query, engine)
    order_data['DateTime'] = pd.to_datetime(order_data['Date'].astype(str) + ' ' + order_data['Time'].astype(str))
    product_data = pd.merge(product_data, order_data, on='Order_ID')

    # Group by Order_ID, Product_ID, Name, and DateTime, then sum the Quantity
    product_data_grouped = product_data.groupby(['Order_ID', 'Product_ID', 'Name', 'DateTime'])['Quantity'].sum().reset_index()
    product_data_grouped.rename(columns={'Quantity': 'Total_Sold'}, inplace=True)

    return product_data_grouped

# Function to generate a daily report
def generate_daily_report():
    data = dataframe.copy()
    today = datetime.now().date()
    daily_data = data[data['DateTime'].dt.date == today]
    return daily_data

# Function to generate a weekly report
def generate_weekly_report():
    data = dataframe.copy()
    today = datetime.now().date()
    last_week = today - timedelta(days=7)
    weekly_data = data[(data['DateTime'].dt.date >= last_week) & (data['DateTime'].dt.date <= today)]
    return weekly_data

# Function to generate a monthly report
def generate_monthly_report():
    data = dataframe.copy()
    today = datetime.now().date()
    last_month = today - timedelta(days=30)
    monthly_data = data[(data['DateTime'].dt.date >= last_month) & (data['DateTime'].dt.date <= today)]
    return monthly_data

def analyze_avg_guest_pax(df, frequency):
    if frequency == 'daily':
        # Define start and end times
        start_time = datetime.now().replace(hour=11, minute=0, second=0, microsecond=0)
        end_time = start_time + timedelta(days=1)

        # Filter data for 11am today to 12am tomorrow
        df_filtered = df[(df['DateTime'] >= start_time) & (df['DateTime'] < end_time)]

        avg_guest_pax_hourly = df_filtered['Guest_Pax'].mean()
        return avg_guest_pax_hourly

    elif frequency == 'weekly':
        # Group by week and calculate average guest pax
        df['Week'] = df['DateTime'].dt.isocalendar().week
        avg_guest_pax_weekly = df.groupby('Week')['Guest_Pax'].mean()
        return avg_guest_pax_weekly

    elif frequency == 'monthly':
        # Group by month and calculate average guest pax
        df['Month'] = df['DateTime'].dt.month
        avg_guest_pax_monthly = df.groupby('Month')['Guest_Pax'].mean()
        return avg_guest_pax_monthly

    else:
        raise ValueError("Invalid frequency. Choose from 'daily', 'weekly', or 'monthly'.")

def analyze_preferred_soup_variations(df, frequency):
    if frequency == 'daily':
        # Define start and end times
        start_time = datetime.now().replace(hour=11, minute=0, second=0, microsecond=0)
        end_time = start_time + timedelta(days=1)

        # Filter data for 11am today to 12am tomorrow
        df_filtered = df[(df['DateTime'] >= start_time) & (df['DateTime'] < end_time)]

        preferred_soup_variations = df_filtered['Soup_Variation'].value_counts()
        return preferred_soup_variations

    elif frequency == 'weekly':
        preferred_soup_variations = df['Soup_Variation'].value_counts()
        return preferred_soup_variations

    elif frequency == 'monthly':
        preferred_soup_variations = df['Soup_Variation'].value_counts()
        return preferred_soup_variations
    else:
        raise ValueError("Invalid frequency. Choose from 'daily', 'weekly', or 'monthly'.")


def analyze_best_selling_product(df, frequency):
    # Load product data
    product_data = load_product_data()

    # Convert DateTime to the correct frequency
    if frequency == 'daily':
        product_data['Period'] = product_data['DateTime'].dt.date
    elif frequency == 'weekly':
        product_data['Period'] = product_data['DateTime'].dt.to_period('W')
    elif frequency == 'monthly':
        product_data['Period'] = product_data['DateTime'].dt.to_period('M')
    else:
        raise ValueError("Invalid frequency. Choose from 'daily', 'weekly', or 'monthly'.")

    # Group by Period, Product_ID and Name, then sum the Total_Sold
    product_data_grouped = product_data.groupby(['Period', 'Product_ID', 'Name'])['Total_Sold'].sum().reset_index()

    # Sort by Total_Sold in descending order and select the top product for each period
    best_selling_product = product_data_grouped.sort_values('Total_Sold', ascending=False).groupby('Period').first().reset_index()

    return best_selling_product

# Save Report to Word
def save_report_to_word(df, frequency, file_path, avg_guest_pax, preferred_soup_variations, best_selling_product):
    document = Document()

    today = datetime.today()
    date_today = today.strftime("%B %d, %Y")
    last_month = today - timedelta(days=30)
    last_month_range = f'{last_month.strftime("%B %d, %Y")} to {today.strftime("%B %d, %Y")}'
    last_week = today - timedelta(days=7)
    last_week_range = f'{last_week.strftime("%B %d, %Y")} to {today.strftime("%B %d, %Y")}'

    # Set up sections and headers
    section = document.sections[0]
    header = section.header
    month_year_today = today.strftime("%B %Y")
    content_header = [
        "Moon Hey Hotpot and Grill",
        "848A Banawe St, Quezon City, 1114 Metro Manila",
        "0917 624 9289",
        f"Trend Analysis {frequency} Report",
        f"({month_year_today})",
    ]

    # Add header content
    for content_h in content_header:
        header_paragraph = header.add_paragraph(content_h)
        header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_paragraph.paragraph_format.space_before = Pt(0)
        header_paragraph.paragraph_format.space_after = Pt(0)

    document.add_heading(f"Average Guest Pax", level=1)

    # Assuming avg_guest_pax is a Series and you want the mean value
    if isinstance(avg_guest_pax, pd.Series):
        avg_guest_pax_value = avg_guest_pax.mean()
    else:
        avg_guest_pax_value = avg_guest_pax  # Assuming avg_guest_pax is already a scalar

    avg_guest_pax_section = document.add_paragraph()
    avg_guest_pax_text = (
        f"The average number of guests per visit during the {frequency.lower()} period was "
        f"{avg_guest_pax_value:.2f}. This indicates that on average, each customer visit "
        f"had approximately {avg_guest_pax_value:.2f} guests."
    )
    avg_guest_pax_section.add_run(avg_guest_pax_text)

    # Add Preferred Soup Variations section
    document.add_heading(f"Preferred Soup Variations", level=1)

    # Define a color dictionary
    color_dict = {
        'Mala soup': '#FF5733',
        'Plain soup': '#FFC300',
        'Suan la soup': '#DAF7A6',
        'Tomato soup': '#C70039'
    }

    # Ensure all soup variations are present in the data
    for soup_variation in color_dict.keys():
        if soup_variation not in preferred_soup_variations.index:
            preferred_soup_variations.loc[soup_variation] = 0

    # Sort the index to ensure consistent order
    preferred_soup_variations.index = preferred_soup_variations.index.astype(str)
    preferred_soup_variations.sort_index(inplace=True)

    # Get the colors for the bars
    bar_colors = [color_dict.get(soup_variation, '#000000') for soup_variation in preferred_soup_variations.index]

    # Plot preferred soup variations
    plt.figure(figsize=(8, 6))
    preferred_soup_variations.plot(kind='bar', color=bar_colors, zorder=2)
    if frequency == 'Daily':
        plt.title(f'Preferred Soup Variations ({date_today})')
    elif frequency == 'Weekly':
        plt.title(f'Preferred Soup Variations ({last_week_range})')
    elif frequency == 'Monthly':
        plt.title(f'Preferred Soup Variations ({last_month_range})')
    plt.xlabel('Soup Variations')
    plt.xticks(rotation=45)
    plt.grid(True, zorder=1)
    plt.ylabel('Count')
    plt.tight_layout()
    soup_path = f'{file_path}/preferred_soup_{frequency.lower()}.png'
    plt.savefig(soup_path)
    plt.close()

    # Insert the plot into the document
    document.add_picture(soup_path, width=Inches(6))

    # Add Best Selling Product section
    document.add_heading(f"Best Selling Product", level=1)
    for period, product in best_selling_product.iterrows():
        document.add_paragraph(f"Period: {product['Period']} - {product['Name']} (Total Sold: {product['Total_Sold']})")

    # Generate and insert the best selling product plot
    best_selling_product_grouped = best_selling_product.groupby('Name')['Total_Sold'].sum()
    plt.figure(figsize=(8, 6))
    best_selling_product_grouped.plot(kind='bar', color='#FF5733', zorder=2)
    if frequency == 'Daily':
        plt.title(f'Best Selling Product ({date_today})')
    elif frequency == 'Weekly':
        plt.title(f'Best Selling Product ({last_week_range})')
    elif frequency == 'Monthly':
        plt.title(f'Best Selling Product ({last_month_range})')
    plt.xlabel('Product')
    plt.xticks(rotation=45)
    plt.grid(True, zorder=1)
    plt.ylabel('Total Sold')
    plt.tight_layout()
    best_selling_product_plot_path = f'{file_path}/best_selling_product_{frequency.lower()}.png'
    plt.savefig(best_selling_product_plot_path)
    plt.close()

    # Insert the best selling product plot into the document
    document.add_picture(best_selling_product_plot_path, width=Inches(6))

    # Add footer
    footer = section.footer
    date = datetime.now().strftime("%B %d, %Y")
    time = datetime.now().strftime("%I:%M %p")
    user_manager = userManager._instance
    username = user_manager.get_current_username()
    footer_paragraph = footer.add_paragraph()
    footer_paragraph.text = f"{date} | {time}    Created by: {username}"
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Save the document as a .docx file
    docx_filename = f'{file_path}/{frequency}_trend_report_{date}.docx'
    document.save(docx_filename)

    # Convert the .docx file to a .pdf file
    pdf_filename = f'{file_path}/{frequency}_trend_report_{date}.pdf'
    convert(docx_filename, pdf_filename)
    print(f"{frequency.capitalize()} report has been converted to PDF and saved to '{pdf_filename}'")

    # Delete the .docx file after conversion
    os.remove(docx_filename)
