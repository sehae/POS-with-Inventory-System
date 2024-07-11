import configparser
import os

import pandas as pd
from datetime import datetime, timedelta

from PyQt5.QtCore import QDateTime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx2pdf import convert
from matplotlib.dates import date2num, DateFormatter
from sqlalchemy import create_engine
import matplotlib.pyplot as plt

from server.create_engine import get_db_engine
from validator.user_manager import userManager


def load_config():
    config = configparser.ConfigParser()
    config.read('config.ini')
    return config


def save_config(report_path):
    config = configparser.ConfigParser()
    config.read('config.ini')

    if report_path:
        if 'DEFAULT' not in config:
            config['DEFAULT'] = {}
        config['DEFAULT']['report_path'] = report_path

    with open('config.ini', 'w') as configfile:
        config.write(configfile)

engine = get_db_engine()

# Query to join product, supplier, and inventory tables
query = """
SELECT 
    p.Product_ID, 
    p.Date,
    p.Time,
    p.Name, 
    p.Original_Quantity, 
    p.Quantity,
    p.Threshold_Value, 
    p.Expiry_Date, 
    p.Availability, 
    p.Category, 
    p.Status,
    i.Selling_Cost,
    i.Buying_Cost,
    COALESCE(s.Supplier_Name, 'Unknown') as Supplier_Name,
    COALESCE(s.Contact_Number, 'Unknown') as Contact_Number,
    COALESCE(s.Email, 'Unknown') as Email,
    COALESCE(s.Address, 'Unknown') as Address,
    COALESCE(s.Status, 'Unknown') as Supplier_Status
FROM 
    product p
LEFT JOIN 
    inventory i ON p.Product_ID = i.Product_ID
LEFT JOIN 
    supplier s ON i.Supplier_ID = s.Supplier_ID
"""


dataframe = pd.read_sql(query, engine)

# Combine Date and Time into a single DateTime column
dataframe['DateTime'] = dataframe['Date'].astype(str) + ' ' + dataframe['Time'].astype(str)

# Function to generate a daily report
def generate_daily_report():
    data = dataframe.copy()
    # Ensure the DateTime column is in the correct format
    data['Date'] = pd.to_datetime(data['DateTime'], errors='coerce')

    # Debugging: Check the date range in the dataset
    print(f"Data date range: {data['Date'].min()} to {data['Date'].max()}")

    # Get today's date
    today = datetime.now().date()
    print(f"Today's date: {today}")

    # Filter data for today
    daily_data = data[data['Date'].dt.date == today]
    print(f"Total records for today ({today}): {len(daily_data)}")

    # Debugging: Print a few records to ensure correctness
    print(daily_data.head())

    return daily_data


# Function to generate a weekly report
def generate_weekly_report():
    data = dataframe.copy()
    data['Date'] = pd.to_datetime(data['DateTime'])
    today = datetime.now().date()
    last_week = today - timedelta(days=7)
    weekly_data = data[(data['Date'] >= pd.to_datetime(last_week)) & (data['Date'] <= pd.to_datetime(today))]
    return weekly_data


# Function to generate a monthly report
def generate_monthly_report():
    data = dataframe.copy()
    data['Date'] = pd.to_datetime(data['DateTime'])
    today = datetime.now().date()
    last_month = today - timedelta(days=30)
    monthly_data = data[(data['Date'] >= pd.to_datetime(last_month)) & (data['Date'] <= pd.to_datetime(today))]
    return monthly_data


# Function to generate plots
def plot_reports(report_data, frequency, file_path):
    today = datetime.now().today()
    date_today = today.strftime('%B %d, %Y')
    last_month = today - timedelta(days=30)
    last_month_range = f'{last_month.strftime("%B %d, %Y")} to {today.strftime("%B %d, %Y")}'
    last_week = today - timedelta(days=7)
    last_week_range = f'{last_week.strftime("%B %d, %Y")} to {today.strftime("%B %d, %Y")}'

    # Inventory Levels by Product
    plt.figure(figsize=(10, 6))
    plt.bar(report_data['Name'], report_data['Quantity'], color='orange', zorder=2)
    plt.xlabel('Product Name')
    plt.ylabel('Quantity')
    plt.title(f'Inventory Levels by Product ({frequency})')

    if frequency == 'Daily':
        plt.title(f'({frequency}) Inventory Levels by Product ({date_today})')
    elif frequency == 'Weekly':
        plt.title(f'({frequency}) Inventory Levels by Product ({last_week_range})')
    else:
        plt.title(f'({frequency}) Inventory Levels by Product ({last_month_range})')

    plt.xticks(rotation=45)
    plt.plot(True)
    plt.tight_layout()
    plt.grid(True, zorder=1)
    plt.savefig(f'{file_path}/inventory_levels_{frequency.lower()}.png')
    plt.close()

    # Inventory Status Overview
    status_counts = report_data['Status'].value_counts()
    plt.figure(figsize=(8, 6))

    plt.pie(status_counts, labels=status_counts.index,
            autopct=lambda pct: f'{pct:.1f}% ({int(pct / 100 * sum(status_counts))})', startangle=140)
    if frequency == 'Daily':
        plt.title(f'Inventory Status Overview ({date_today})')
    elif frequency == 'Weekly':
        plt.title(f'Inventory Status Overview ({last_week_range})')
    else:
        plt.title(f'Inventory Status Overview ({last_month_range})')

    plt.tight_layout()
    plt.savefig(f'{file_path}/inventory_status_{frequency.lower()}.png')
    plt.close()

    # Expiry Date Analysis
    report_data['Expiry_Date'] = pd.to_datetime(report_data['Expiry_Date'])
    upcoming_expiry = report_data[report_data['Expiry_Date'] >= datetime.now()].copy()
    upcoming_expiry['Expiry_Date'] = upcoming_expiry['Expiry_Date'].dt.date  # Convert to date for clarity
    expiry_counts = upcoming_expiry['Expiry_Date'].value_counts().sort_index()

    expiry_dates_num = date2num(expiry_counts.index)

    # Plotting
    plt.figure(figsize=(12, 6))
    plt.plot(expiry_dates_num, expiry_counts.values, marker='o', linestyle='-')
    plt.xlabel('Expiry Date')
    plt.ylabel('Number of Products')
    if frequency == 'Daily':
        plt.title(f'{frequency} Upcoming Expiry Date Analysis ({date_today})')
    elif frequency == 'Weekly':
        plt.title(f'{frequency} Upcoming Expiry Date Analysis ({last_week_range})')
    else:
        plt.title(f'{frequency} Upcoming Expiry Date Analysis ({last_month_range})')

    # Format the x-ticks as requested
    date_format = DateFormatter('%B %d, %Y')
    plt.gca().xaxis.set_major_formatter(date_format)

    plt.xticks(rotation=45)
    plt.tight_layout()
    plt.grid(True)
    plt.savefig(f'{file_path}/expiry_date_time_series_{frequency.lower()}.png')
    plt.close()


def save_report_to_excel(report_data, report_type, file_path):
    filename = f'{file_path}/{report_type}_report.xlsx'
    report_data.to_excel(filename, index=False)
    print(f"{report_type.capitalize()} report has been generated and saved to '{filename}'")


def save_report_to_word(report_data, report_type, file_path):
    document = Document()

    # Set up sections and headers
    section = document.sections[0]
    header = section.header
    today = datetime.today()
    month_year_today = today.strftime("%B %Y")
    content_header = [
        "Moon Hey Hotpot and Grill",
        "848A Banawe St, Quezon City, 1114 Metro Manila",
        "0917 624 9289",
        f"Inventory {report_type} Report",
        f"({month_year_today})",
    ]
    for content_h in content_header:
        header_paragraph = header.add_paragraph(content_h)
        run = header_paragraph.runs[0]

        # Center align the paragraph
        header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Remove spacing before and after the paragraph
        header_paragraph.paragraph_format.space_before = Pt(0)
        header_paragraph.paragraph_format.space_after = Pt(0)

    footer = section.footer
    date = datetime.now().strftime("%B %d, %Y")
    time = datetime.now().strftime("%I:%M %p")
    user_manager = userManager._instance
    username = user_manager.get_current_username()
    footer_paragraph = footer.add_paragraph()
    footer_paragraph.text = f"{date} | {time}    Created by: {username}"
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Compute total spending per supplier
    suppliers = report_data['Supplier_Name'].unique()
    supplier_totals = {}
    for supplier in suppliers:
        supplier_data = report_data[report_data['Supplier_Name'] == supplier]
        total_spent = (supplier_data['Buying_Cost'] * supplier_data['Original_Quantity']).sum()
        supplier_totals[supplier] = total_spent

    # Add supplier-wise information
    for supplier, total_spent in supplier_totals.items():
        document.add_heading(f"Supplier: {supplier}", level=1)
        document.add_paragraph(f"Total spent with this supplier: ₱{total_spent:,.2f}\n")

        # List products and spending per product
        supplier_data = report_data[report_data['Supplier_Name'] == supplier]
        products = supplier_data['Name'].unique()
        for product in products:
            product_data = supplier_data[supplier_data['Name'] == product]
            total_product_spent = (product_data['Buying_Cost'] * product_data['Original_Quantity']).sum()
            price = product_data['Buying_Cost'].iloc[0]

            # Include expiry date and date of purchase
            expiry_date = product_data['Expiry_Date'].iloc[0].strftime('%B %d, %Y')
            date_of_purchase = product_data['Date'].iloc[0].strftime('%B %d, %Y')

            document.add_paragraph(f"Product: {product}")
            document.add_paragraph(f"Total spent on this product: ₱{total_product_spent:,.2f}")
            document.add_paragraph(f"Price per unit: ₱{price:,.2f}")
            document.add_paragraph(f"Quantity bought: {product_data['Original_Quantity'].sum()}")
            document.add_paragraph(f"Expiry Date: {expiry_date}")
            document.add_paragraph(f"Date of Purchase: {date_of_purchase}\n")

        # Total spending with this supplier for all products
        document.add_paragraph(f"Total spent on all products with {supplier}: ₱{total_spent:,.2f}\n")

    # Add plots to the Word document
    plot_paths = {
        'Inventory Levels': f'{file_path}/inventory_levels_{report_type.lower()}.png',
        'Inventory Status': f'{file_path}/inventory_status_{report_type.lower()}.png',
        'Expiry Date Analysis': f'{file_path}/expiry_date_time_series_{report_type.lower()}.png'
    }

    for title, path in plot_paths.items():
        document.add_heading(f"{title} ({report_type})", level=1)
        document.add_picture(path, width=Inches(6))

    # Save the document as a .docx file
    docx_filename = f'{file_path}/{report_type}_Inventory_report_{date}.docx'
    document.save(docx_filename)

    # Convert the .docx file to a .pdf file
    pdf_filename = f'{file_path}/{report_type}_Inventory_report_{date}.pdf'
    convert(docx_filename, pdf_filename)
    print(f"{report_type.capitalize()} report has been converted to PDF and saved to '{pdf_filename}'")

    # Delete the .docx file after conversion
    os.remove(docx_filename)
