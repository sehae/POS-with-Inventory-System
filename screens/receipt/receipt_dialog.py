import sys
from PyQt5.QtWidgets import QDialog, QVBoxLayout, QLabel, QPushButton, QHBoxLayout, QMessageBox
from docx import Document
from docx.shared import Inches
import os

class ReceiptDialog(QDialog):
    def __init__(self, order_details):
        super().__init__()
        self.setWindowTitle("Order Receipt")
        self.order_details = order_details

        main_layout = QVBoxLayout()

        # Example: Display order details
        details_label = QLabel(order_details)
        main_layout.addWidget(details_label)

        # Button layout
        button_layout = QHBoxLayout()
        print_button = QPushButton("Print Now")
        cancel_button = QPushButton("Cancel")
        button_layout.addWidget(print_button)
        button_layout.addWidget(cancel_button)

        # Connect buttons to slots or functions
        print_button.clicked.connect(self.print_order)
        cancel_button.clicked.connect(self.reject)  # Close the dialog on cancel

        main_layout.addLayout(button_layout)
        self.setLayout(main_layout)

    def print_order(self):
        try:
            # Create a Word document
            document = Document()

            # Set the page width to 80mm and height to a reasonable size
            section = document.sections[0]
            section.page_width = Inches(3.14961)  # 80mm to inches
            section.page_height = Inches(11.69)  # Arbitrary height, can be adjusted

            # Remove the margins
            section.left_margin = Inches(0.19685)
            section.right_margin = Inches(0.19685)
            section.top_margin = Inches(0.19685)
            section.bottom_margin = Inches(0)

            # Add header
            header_paragraph = document.add_paragraph()
            header_paragraph.alignment = 1  # Center alignment
            header_paragraph.add_run("MOON HEY HOTPOT AND GRILL\n848A Banawe St, Quezon City, 1114 Metro Manila\nContact Number: 0917 123 4567").bold = True

            # Add order details
            details_lines = self.order_details.split('\n')[3:]  # Skipping the first three lines which are used in the header
            for line in details_lines:
                if ':' in line:
                    label, value = line.split(':', 1)
                    p = document.add_paragraph()
                    p.add_run(label.strip() + ': ').bold = True
                    p.add_run(value.strip())
                else:
                    p = document.add_paragraph(line)
                    p.alignment = 1  # Center alignment

            # Add thank you note
            thank_you_paragraph = document.add_paragraph()
            thank_you_paragraph.alignment = 1  # Center alignment
            thank_you_paragraph.add_run("Thank you for your visit!")

            # Save the document
            file_path = os.path.join(os.getcwd(), 'receipt.docx')
            document.save(file_path)

            # Open the document for printing (platform-dependent)
            if sys.platform == "win32":
                os.startfile(file_path, "print")
            elif sys.platform == "darwin":
                os.system(f"open -a 'Preview' {file_path}")
            else:
                os.system(f"libreoffice --headless --print-to-file --printer-name=PrinterName {file_path}")

            # Inform the user that the receipt has been generated
            QMessageBox.information(self, "Receipt Generated", "The receipt has been generated and sent to the printer.")
            self.accept()  # Close the dialog after printing

        except FileNotFoundError:
            QMessageBox.critical(self, "Error", "Printer not found.")

        except PermissionError:
            QMessageBox.critical(self, "Error", "Permission denied to access printer.")

        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error printing: {str(e)}")
