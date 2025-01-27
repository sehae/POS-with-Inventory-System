import sys
from PyQt5 import QtWidgets, QtCore
from PyQt5.QtWidgets import QDialog, QVBoxLayout, QLabel, QPushButton, QHBoxLayout, QGridLayout, QMessageBox
from docx import Document
from docx.shared import Inches
import os

class CheckoutReceiptDialog(QDialog):
    def __init__(self, order_details, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Order Receipt")
        self.order_details = order_details

        main_layout = QVBoxLayout()

        # Header details
        header_label = QLabel("MOON HEY HOTPOT AND GRILL\n848A Banawe St, Quezon City, 1114 Metro Manila\nContact Number: 0917 123 4567")
        header_label.setAlignment(QtCore.Qt.AlignCenter)
        main_layout.addWidget(header_label)

        # Creating a grid layout for the order details
        grid_layout = QGridLayout()
        grid_layout.setHorizontalSpacing(10)  # Set spacing between columns

        # Splitting order details into lines for label-value pairs
        details_lines = order_details.split('\n')[3:]  # Skipping the first three lines which are used in the header
        for i, line in enumerate(details_lines):
            if ':' in line:  # Check if line is in label:value format
                label, value = line.split(':', 1)  # Split only on the first colon
                label_widget = QLabel(label.strip() + ':')
                value_widget = QLabel(value.strip())
                label_widget.setAlignment(QtCore.Qt.AlignRight | QtCore.Qt.AlignVCenter)
                value_widget.setAlignment(QtCore.Qt.AlignLeft | QtCore.Qt.AlignVCenter)
                grid_layout.addWidget(label_widget, i, 0)
                grid_layout.addWidget(value_widget, i, 1)
            else:  # If not, assume it's a standalone line
                single_line_label = QLabel(line)
                single_line_label.setAlignment(QtCore.Qt.AlignCenter)
                grid_layout.addWidget(single_line_label, i, 0, 1, 2)  # Span both columns

        main_layout.addLayout(grid_layout)

        # Thank you label
        thank_you_label = QLabel("Thank you for your visit!")
        thank_you_label.setAlignment(QtCore.Qt.AlignCenter)
        main_layout.addWidget(thank_you_label)

        # Button layout
        button_layout = QHBoxLayout()
        print_button = QPushButton("Print Now")
        cancel_button = QPushButton("Cancel")
        button_layout.addWidget(print_button)
        button_layout.addWidget(cancel_button)

        # Connect buttons to slots or functions
        print_button.clicked.connect(self.print_order)
        cancel_button.clicked.connect(self.reject)  # Close the dialog on cancel

        main_layout.addLayout(button_layout)
        self.setLayout(main_layout)

        self.adjustSize()  # Adjust size based on content
        self.centerDialog()

    def centerDialog(self):
        # Center dialog within its parent or screen
        frame_geometry = self.frameGeometry()
        screen_center = QtWidgets.QDesktopWidget().availableGeometry().center()
        frame_geometry.moveCenter(screen_center)
        self.move(frame_geometry.topLeft())

    def print_order(self):
        # Create a Word document
        document = Document()

        # Set the page width to 80mm and height to a reasonable size
        section = document.sections[0]
        section.page_width = Inches(3.14961)  # 80mm to inches
        section.page_height = Inches(11.69)  # Arbitrary height, can be adjusted

        # Remove the margins
        section.left_margin = Inches(0.19685)
        section.right_margin = Inches(0.19685)
        section.top_margin = Inches(0.19685)
        section.bottom_margin = Inches(0)

        # Add header
        header_paragraph = document.add_paragraph()
        header_paragraph.alignment = 1  # Center alignment
        header_paragraph.add_run("MOON HEY HOTPOT AND GRILL\n848A Banawe St, Quezon City, 1114 Metro Manila\nContact Number: 0917 123 4567").bold = True

        # Add order details
        details_lines = self.order_details.split('\n')[3:]  # Skipping the first three lines which are used in the header
        for line in details_lines:
            if ':' in line:
                label, value = line.split(':', 1)
                p = document.add_paragraph()
                p.add_run(label.strip() + ': ').bold = True
                p.add_run(value.strip())
            else:
                p = document.add_paragraph(line)
                p.alignment = 1  # Center alignment

        # Add thank you note
        thank_you_paragraph = document.add_paragraph()
        thank_you_paragraph.alignment = 1  # Center alignment
        thank_you_paragraph.add_run("Thank you for your visit!")

        # Save the document
        file_path = os.path.join(os.getcwd(), 'receipt.docx')
        document.save(file_path)

        # Open the document for printing (platform-dependent)
        if sys.platform == "win32":
            os.startfile(file_path, "print")
        elif sys.platform == "darwin":
            os.system(f"open -a 'Preview' {file_path}")
        else:
            os.system(f"libreoffice --headless --print-to-file --printer-name=PrinterName {file_path}")

        # Inform the user that the receipt has been generated
        QMessageBox.information(self, "Receipt Generated", "The receipt has been generated and sent to the printer.")